# src/utils.py
"""
Utility functions for EduPlan PH.
Handles document export (DOCX, PDF) and data export (CSV).
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fpdf import FPDF
import pandas as pd


def sanitize_for_pdf(text: str) -> str:
    """Sanitize text for PDF rendering with Helvetica (latin-1 only)."""
    replacements = {
        '\u2013': '-',   # en dash
        '\u2014': '-',   # em dash
        '\u2015': '-',   # horizontal bar
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2026': '...',  # ellipsis
        '\u2022': '-',   # bullet
        '\u00a0': ' ',   # non-breaking space
        '\u2010': '-',   # hyphen
        '\u2011': '-',   # non-breaking hyphen
        '\u2012': '-',   # figure dash
        '\u00b7': '-',   # middle dot
        '\u2027': '-',   # hyphenation point
        '\u25cf': '-',   # black circle
        '\u25cb': 'o',   # white circle
        '\u2192': '->',  # right arrow
        '\u2190': '<-',  # left arrow
        '\u2714': '[/]',  # check mark
        '\u2716': '[x]',  # cross mark
        '\u00f1': 'n',   # ñ - common in Filipino
        '\u00d1': 'N',   # Ñ
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Remove any remaining non-latin-1 characters
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    return text


def markdown_to_plain_text(md_content: str) -> str:
    """Convert basic Markdown to plain text for document export."""
    text = md_content
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'---+', '', text)
    return text.strip()


def parse_markdown_lines(md_content: str) -> list:
    """
    Parse markdown content into structured lines with type info.
    Returns list of dicts: {'type': str, 'text': str, 'level': int}
    Types: 'h1', 'h2', 'h3', 'section_header', 'sub_header', 'bullet', 'numbered', 'text', 'empty'
    """
    lines = []
    for raw_line in md_content.split('\n'):
        line = raw_line.rstrip()

        if not line.strip():
            lines.append({'type': 'empty', 'text': '', 'level': 0})
            continue

        # Markdown headings
        h_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            # Strip bold markers from headings
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            lines.append({'type': f'h{min(level, 3)}', 'text': text, 'level': level})
            continue

        stripped = line.strip()

        # Roman numeral section headers (I. OBJECTIVES, II. CONTENT, etc.)
        if re.match(r'^[IVX]+\.\s', stripped):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            lines.append({'type': 'section_header', 'text': text, 'level': 1})
            continue

        # Letter sub-headers (A. Preliminary Activities, B. Lesson Proper, etc.)
        if re.match(r'^[A-H]\.\s', stripped):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            lines.append({'type': 'sub_header', 'text': text, 'level': 2})
            continue

        # Detect indentation level
        indent = len(raw_line) - len(raw_line.lstrip())
        indent_level = indent // 3  # roughly 3 spaces per level

        # Bullet points (-, *, •)
        bullet_match = re.match(r'^(\s*)[*\-•]\s+(.*)', line)
        if bullet_match:
            text = bullet_match.group(2).strip()
            # Preserve bold in bullets
            lines.append({'type': 'bullet', 'text': text, 'level': indent_level})
            continue

        # Numbered items
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            text = num_match.group(2).strip()
            lines.append({'type': 'numbered', 'text': text, 'level': indent_level})
            continue

        # Regular text
        text = stripped
        lines.append({'type': 'text', 'text': text, 'level': indent_level})

    return lines


def export_to_docx(content: str, topic: str, grade_level: str, subject: str) -> bytes:
    """
    Generate a properly formatted Word document from the lesson plan content.
    Produces a clean, professional DepEd DLP format.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(2)

    # Title
    title = doc.add_heading('DETAILED LESSON PLAN', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(27, 42, 74)  # Deep Navy

    # Meta info
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(4)
    meta_run = meta_para.add_run(
        f"Subject: {subject}  |  Grade Level: {grade_level}  |  Topic: {topic}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    # Thin line separator
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse and format content
    parsed = parse_markdown_lines(content)

    for item in parsed:
        if item['type'] == 'empty':
            continue

        text = item['text']
        # Clean markdown bold markers for display
        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

        if item['type'] in ('h1', 'section_header'):
            para = doc.add_heading(clean_text, level=1)
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(27, 42, 74)  # Deep Navy

        elif item['type'] in ('h2', 'sub_header'):
            para = doc.add_heading(clean_text, level=2)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(30, 30, 30)

        elif item['type'] == 'h3':
            para = doc.add_heading(clean_text, level=3)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)

        elif item['type'] == 'bullet':
            indent_cm = 1.0 + (item['level'] * 0.8)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(indent_cm)
            para.paragraph_format.first_line_indent = Cm(-0.4)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(1)

            # Add bullet character + text with bold handling
            _add_formatted_run(para, '\u2022  ', text, font_size=Pt(11))

        elif item['type'] == 'numbered':
            indent_cm = 1.0 + (item['level'] * 0.8)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(indent_cm)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(1)

            _add_formatted_run(para, '', text, font_size=Pt(11))

        else:
            # Regular text
            indent_cm = max(0, item['level'] * 0.5)
            para = doc.add_paragraph()
            if indent_cm > 0:
                para.paragraph_format.left_indent = Cm(indent_cm)
            para.paragraph_format.space_after = Pt(3)

            _add_formatted_run(para, '', text, font_size=Pt(11))

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Generated by EduPlan PH | AI-Enhanced Lesson Plan Generator"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_formatted_run(para, prefix: str, text: str, font_size=None):
    """
    Add text to a paragraph, preserving **bold** markdown markers as actual bold formatting.
    """
    if prefix:
        run = para.add_run(prefix)
        run.font.name = 'Arial'
        if font_size:
            run.font.size = font_size

    # Split text by bold markers and alternate between normal and bold
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        else:
            run = para.add_run(part)
        run.font.name = 'Arial'
        if font_size:
            run.font.size = font_size


def export_to_pdf(content: str, topic: str, grade_level: str, subject: str) -> bytes:
    """
    Generate a PDF document from the lesson plan content.
    Matches the DOCX alignment and formatting exactly.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    L_MARGIN = 15
    R_MARGIN = 15
    pdf.set_margins(L_MARGIN, 15, R_MARGIN)
    pdf.add_page()
    
    usable_w = pdf.w - L_MARGIN - R_MARGIN
    
    # ---- Title Block ----
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, sanitize_for_pdf("DETAILED LESSON PLAN"), ln=True, align="C")

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    meta = f"Subject: {subject}  |  Grade Level: {grade_level}  |  Topic: {topic}"
    pdf.cell(0, 7, sanitize_for_pdf(meta), ln=True, align="C")
    pdf.set_text_color(0, 0, 0)

    # Separator
    pdf.set_draw_color(180, 180, 180)
    pdf.line(L_MARGIN, pdf.get_y() + 1, pdf.w - R_MARGIN, pdf.get_y() + 1)
    pdf.ln(5)

    # ---- Parse the markdown content ----
    parsed = parse_markdown_lines(content)

    for item in parsed:
        itype = item['type']
        raw_text = item['text']
        level = item['level']

        if itype == 'empty':
            pdf.ln(2)
            continue
            
        if raw_text.strip() in ('---', '----', '-----', '------'):
            y = pdf.get_y()
            pdf.set_draw_color(200, 200, 200)
            pdf.line(L_MARGIN, y, pdf.w - R_MARGIN, y)
            pdf.ln(3)
            continue

        safe_text = sanitize_for_pdf(raw_text)
        
        # Reset left margin for this line
        current_indent = 0

        # ---- Section headers: I. OBJECTIVES, II. CONTENT, etc. ----
        if itype in ('h1', 'section_header'):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', safe_text)
            pdf.ln(4)
            pdf.set_left_margin(L_MARGIN + 3)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(27, 42, 74)  # Deep Navy
            # Draw line properly based on string width
            w = pdf.get_string_width(clean) + 2
            pdf.multi_cell(0, 7, clean)
            
            y = pdf.get_y()
            pdf.set_draw_color(184, 151, 59)  # Muted Gold accent line
            pdf.line(L_MARGIN + 3, y, L_MARGIN + 3 + min(w, 80), y)
            pdf.set_draw_color(0, 0, 0)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)

        # ---- Sub headers: A. Preliminary Activities, B. Lesson Proper, etc. ----
        elif itype in ('h2', 'sub_header'):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', safe_text)
            pdf.ln(3)
            pdf.set_left_margin(L_MARGIN + 6)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, clean)
            pdf.ln(1)

        # ---- h3 sub-sub headers ----
        elif itype == 'h3':
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', safe_text)
            pdf.set_left_margin(L_MARGIN + 10)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 5.5, clean)

        # ---- Bullet points ----
        elif itype == 'bullet':
            current_indent = 10 + min(level, 3) * 5
            pdf.set_left_margin(L_MARGIN + current_indent)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "", 10)
            txt = "- " + safe_text
            try:
                pdf.multi_cell(0, 5, txt, markdown=True)
            except Exception:
                # Fallback if markdown parser fails
                pdf.multi_cell(0, 5, re.sub(r'\*\*(.+?)\*\*', r'\1', txt))

        # ---- Numbered items ----
        elif itype == 'numbered':
            current_indent = 10 + min(level, 3) * 5
            pdf.set_left_margin(L_MARGIN + current_indent)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "", 10)
            try:
                pdf.multi_cell(0, 5, safe_text, markdown=True)
            except Exception:
                pdf.multi_cell(0, 5, re.sub(r'\*\*(.+?)\*\*', r'\1', safe_text))

        # ---- Regular text / paragraphs ----
        else:
            current_indent = min(level, 3) * 5
            pdf.set_left_margin(L_MARGIN + current_indent)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "", 10)
            try:
                pdf.multi_cell(0, 5, safe_text, markdown=True)
            except Exception:
                pdf.multi_cell(0, 5, re.sub(r'\*\*(.+?)\*\*', r'\1', safe_text))
                
        # Reset margin back to normal after the block
        pdf.set_left_margin(L_MARGIN)
        pdf.set_x(pdf.l_margin)

    # ---- Footer ----
    pdf.ln(8)
    pdf.set_left_margin(L_MARGIN)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, "Generated by EduPlan PH | AI-Enhanced Lesson Plan Generator", align="C")

    return bytes(pdf.output(dest='S'))


def export_quiz_to_csv(quiz_data: list) -> str:
    """
    Convert extracted quiz data to CSV format for download.
    """
    if not quiz_data:
        return "No quiz data available."

    df = pd.DataFrame(quiz_data)
    df.columns = ["Item #", "Question", "A", "B", "C", "D", "Correct Answer"]
    return df.to_csv(index=False)